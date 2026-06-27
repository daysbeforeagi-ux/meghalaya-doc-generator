"""
Sources Dossier writer (§12.2).
Separate file — never merged with the deliverable.
Keyed to the draft via claim_ids; includes review checklist.
"""
from __future__ import annotations

from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from src.intake.models import DossierEntry, GateCheckResult

_TIER_LABELS = {
    1: "Tier 1 — Official Government",
    2: "Tier 2 — Established Press / Institution",
    3: "Tier 3 — Contextual Background",
}
_CONFIDENCE_LABELS = {
    "high":    "HIGH CONFIDENCE",
    "medium":  "MEDIUM CONFIDENCE — corroborate if critical",
    "flagged": "FLAGGED — requires human review before use",
}


def write_dossier_docx(
    entries: List[DossierEntry],
    gate: GateCheckResult,
    omitted_facts: List[str],
    sensitive_flags: List[str],
    output_path: Path,
) -> None:
    """Write the Sources Dossier to a separate .docx (§12.2)."""
    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

    doc.add_heading("Sources Dossier", level=1)
    intro = doc.add_paragraph(
        "This document is the verification record for the accompanying speech draft. "
        "It is delivered alongside but is NEVER merged with the deliverable. "
        "Every factual claim in the speech is keyed here by claim_id. "
        "Review all items before approving the speech for publication."
    )
    intro.paragraph_format.space_after = Pt(12)

    # ── Verified sources ──────────────────────────────────────────────────────
    doc.add_heading("Verified Sources", level=2)

    if entries:
        for entry in entries:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)

            id_run = p.add_run(f"[{entry.claim_id}]  ")
            id_run.bold = True
            id_run.font.size = Pt(11)

            p.add_run(f"Claim:\n      {entry.claim}\n").font.size = Pt(11)
            p.add_run(f"Source:\n      {entry.publisher} — {entry.source_url}\n").font.size = Pt(10)
            p.add_run(f"Tier:\n      {_TIER_LABELS.get(entry.source_tier, str(entry.source_tier))}\n").font.size = Pt(10)
            p.add_run(f'Evidence:\n      "{entry.evidence_passage}"\n').font.size = Pt(10)
            p.add_run(f"Accessed:   {entry.accessed_at[:10]}\n").font.size = Pt(10)

            conf_label = _CONFIDENCE_LABELS.get(entry.confidence, entry.confidence)
            conf_run = p.add_run(f"Confidence: {conf_label}")
            conf_run.font.size = Pt(10)
            if entry.confidence == "flagged":
                conf_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                conf_run.bold = True
            elif entry.confidence == "medium":
                conf_run.font.color.rgb = RGBColor(0xCC, 0x80, 0x00)
    else:
        doc.add_paragraph("No verified evidence records were produced for this draft.")

    # ── Review checklist (§12.2) ──────────────────────────────────────────────
    doc.add_heading("Review Checklist", level=2)

    checklist_items: list[tuple[str, bool]] = []  # (text, is_problem)

    if omitted_facts:
        doc.add_heading("Omitted / Unsourced Facts (§8.4)", level=3)
        doc.add_paragraph(
            "The following facts could not be verified and were OMITTED from the speech. "
            "If these facts are essential, provide verified sources and regenerate."
        )
        for fact in omitted_facts:
            p = doc.add_paragraph(f"• {fact}", style="List Bullet")
            p.runs[0].font.color.rgb = RGBColor(0xCC, 0x60, 0x00)

    if sensitive_flags:
        doc.add_heading("Sensitive Content — Mandatory Human Sign-off Required (§10.5)", level=3)
        for flag in sensitive_flags:
            p = doc.add_paragraph(f"• {flag}", style="List Bullet")
            p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            p.runs[0].bold = True

    flagged_entries = [e for e in entries if e.confidence == "flagged"]
    if flagged_entries:
        doc.add_heading("Flagged-Confidence Claims (§10.1)", level=3)
        doc.add_paragraph(
            "These claims were included but are flagged for low confidence. "
            "Verify independently before publication."
        )
        for e in flagged_entries:
            doc.add_paragraph(f"• [{e.claim_id}] {e.claim}", style="List Bullet")

    if gate.unmapped_claims:
        doc.add_heading("Gate Check Failures", level=3)
        doc.add_paragraph(
            "The following claim_ids appear in the dossier but have no matching evidence record. "
            "This indicates a generator error. Do not publish until resolved."
        )
        for cid in gate.unmapped_claims:
            p = doc.add_paragraph(f"• {cid}")
            p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            p.runs[0].bold = True

    if gate.warnings:
        doc.add_heading("Other Warnings", level=3)
        for w in gate.warnings:
            doc.add_paragraph(f"• {w}", style="List Bullet")

    # ── Gate check summary ────────────────────────────────────────────────────
    doc.add_heading("Gate Check Summary (§13)", level=2)
    status_text = "PASSED ✓" if gate.passed else "FAILED ✗ — review required before publication"
    status_para = doc.add_paragraph(f"Automated gate check: {status_text}")
    status_para.runs[0].bold = True
    if not gate.passed:
        status_para.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # ── Publication footer ────────────────────────────────────────────────────
    doc.add_heading("Publication Status", level=2)
    footer = doc.add_paragraph(
        "⚠  This speech is a MACHINE DRAFT. All factual claims should be independently "
        "verified by a human reviewer before official use. The Sources Dossier provides "
        "traceability but is not a substitute for editorial judgement and political review."
    )
    footer.runs[0].bold = True
    footer.paragraph_format.space_before = Pt(12)

    doc.save(str(output_path))
