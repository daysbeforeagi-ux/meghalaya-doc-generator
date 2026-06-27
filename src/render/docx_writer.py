"""
Deliverable .docx writer (§12.1).
Outputs clean prose with no inline citations, no sentinels, no footnotes.
Presents as a draft for human approval (§10.10).
"""
from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


def write_speech_docx(
    clean_text: str,
    output_path: Path,
    speaker_name: Optional[str] = None,
) -> None:
    """Write clean speech text to a presentation-ready .docx (§12.1)."""
    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

    # Title
    heading_text = f"Speech — {speaker_name}" if speaker_name else "Official Speech"
    title = doc.add_heading(heading_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Draft disclaimer — §10.10
    banner = doc.add_paragraph(
        "DRAFT FOR HUMAN REVIEW — Not for publication without authorised approval."
    )
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in banner.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Speech body — split on double newline for natural paragraph breaks
    paragraphs = clean_text.split("\n\n")
    for para_text in paragraphs:
        lines = [ln.strip() for ln in para_text.split("\n") if ln.strip()]
        for line in lines:
            p = doc.add_paragraph(line)
            # Generous delivery spacing (§12.1)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(22)
            if p.runs:
                p.runs[0].font.size = Pt(13)

    doc.save(str(output_path))
