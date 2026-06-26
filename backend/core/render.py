"""DOCX rendering — clean deliverable + separate Sources Dossier."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

from models.session import EvidenceRecord, Session


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_para(doc: Document, text: str, size: int = 12, bold: bool = False,
               alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def build_deliverable(session: Session, draft: str, output_path: Path) -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Header
    is_speech = session.doc_type and session.doc_type == "speech"
    label = "DRAFT SPEECH" if is_speech else "DRAFT PRESS RELEASE"
    _add_para(doc, label, size=9, bold=True)
    _add_para(doc, f"Prepared: {datetime.now(timezone.utc).strftime('%d %B %Y')}", size=9)
    _add_para(doc, "FOR HUMAN REVIEW BEFORE USE — this is a machine-generated draft.", size=9, bold=True)
    doc.add_paragraph()

    if is_speech and session.speaker:
        role = session.speaker.role.replace("_", " ").title()
        label_line = session.speaker.name if session.speaker.name else role
        _add_para(doc, label_line.upper(), size=14, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

    # Body — each paragraph
    for para_text in draft.split("\n\n"):
        stripped = para_text.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        else:
            p = doc.add_paragraph()
            if is_speech:
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = Pt(24)
            run = p.add_run(stripped)
            run.font.size = Pt(13 if is_speech else 11)

    doc.add_paragraph()
    _add_para(doc, "— END OF DRAFT —", size=9, bold=False,
               alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(str(output_path))


def build_dossier(
    session: Session,
    draft: str,
    claim_map: list[dict],
    evidence: list[EvidenceRecord],
    output_path: Path,
) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    doc.add_heading("SOURCES DOSSIER", level=1)
    _add_para(doc, f"Session: {session.session_id}", size=9)
    _add_para(doc, f"Generated: {datetime.now(timezone.utc).strftime('%d %B %Y, %H:%M UTC')}", size=9)
    _add_para(doc, f"Document type: {session.doc_type}", size=9)
    doc.add_paragraph()

    # Evidence records
    doc.add_heading("Evidence Records", level=2)
    ev_by_id = {ev.claim_id: ev for ev in evidence}

    if evidence:
        for ev in evidence:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{ev.claim_id}]  ").bold = True
            run = p.add_run(f"Claim: {ev.claim}\n")
            run.font.size = Pt(11)
            p.add_run(f"Source:    {ev.publisher} — {ev.source_url}\n")
            p.add_run(f"Tier:      {ev.source_tier} ({'official' if ev.source_tier == 1 else 'press/institution' if ev.source_tier == 2 else 'context'})\n")
            p.add_run(f"Evidence:  \"{ev.supporting_passage}\"\n")
            p.add_run(f"Accessed:  {ev.accessed_at.strftime('%Y-%m-%d')}\n")
            p.add_run(f"Confidence: {ev.confidence}")
            doc.add_paragraph()
    else:
        _add_para(doc, "No external evidence was gathered for this draft.", size=11)

    doc.add_paragraph()

    # Review checklist
    doc.add_heading("Review Checklist", level=2)
    flagged = [ev for ev in evidence if ev.confidence == "flagged"]
    medium = [ev for ev in evidence if ev.confidence == "medium"]

    checklist_items = [
        ("[ ]", "Verify draft is a faithful reflection of the brief"),
        ("[ ]", "Confirm all facts against the Evidence Records above"),
        ("[ ]", "Check that no inline citations appear in the deliverable"),
        ("[ ]", "Confirm voice matches the speaker / register"),
    ]
    if flagged:
        for ev in flagged:
            checklist_items.append(("[ ] FLAGGED", f"[{ev.claim_id}] {ev.claim} — needs verification"))
    if medium:
        for ev in medium:
            checklist_items.append(("[ ] REVIEW", f"[{ev.claim_id}] {ev.claim} — medium confidence"))

    checklist_items += [
        ("[ ]", "Check for sensitive content (religious, communal, caste) — route for human sign-off if present"),
        ("[ ]", "Check for any legal risks (defamation, incitement)"),
        ("[ ]", "Approved by authorised reviewer before use"),
    ]

    for marker, text in checklist_items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{marker}  ").bold = ("FLAGGED" in marker or "REVIEW" in marker)
        p.add_run(text)

    doc.add_paragraph()
    _add_para(
        doc,
        "This dossier is confidential and intended solely for the authorised reviewer. "
        "It must not be distributed with the deliverable.",
        size=9,
    )

    doc.save(str(output_path))
